from fastapi import APIRouter, HTTPException, Body
from typing import Dict, Any, List
import logging
from fastapi.responses import FileResponse
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from docx import Document
from reportlab.pdfgen import canvas
from docx.enum.text import WD_ALIGN_PARAGRAPH
from backend.config.settings import get_settings
import tempfile
import os
from pydantic import BaseModel

router = APIRouter(prefix="/drafter", tags=["drafter"])
settings = get_settings()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize the model
model = ChatGroq(
    api_key=settings.API_KEY,
    model_name=settings.MODEL_NAME,
    temperature=settings.TEMPERATURE
)

# Define request models
class DocumentRequest(BaseModel):
    case_details: str
    ipc_sections: List[str]

class EditRequest(BaseModel):
    document_content: str

@router.post("/generate_document")
async def generate_document(request: DocumentRequest) -> Dict[str, str]:
    try:
        logger.info("Generating legal document")
        # Create prompt template
        prompt = PromptTemplate.from_template(
            """Create a formal legal document that follows standard legal formatting. Format the response in proper HTML for a clean layout.

Case Details: {case_details}
IPC Sections: {ipc_sections}

Format the document with the following structure:
<div class='legal-document'>
    <div class='header' style='text-align: center; margin-bottom: 20px;'>
        <h1>IN THE COURT OF [COURT NAME], [STATE]</h1>
        <p>Criminal Original Jurisdiction</p>
        <p>Case No. [Year]</p>
    </div>
    
    <div class='parties' style='margin-bottom: 20px;'>
        <p>State of [State Name]</p>
        <p style='text-align: right;'>...Complainant</p>
        <p style='text-align: center;'>Versus</p>
        <p>[Accused Name]</p>
        <p style='text-align: right;'>...Accused</p>
    </div>
    
    <div class='application-title' style='margin-bottom: 20px;'>
        <p style='text-align: center;'><strong>APPLICATION UNDER SECTION 302 OF THE INDIAN PENAL CODE, 1860</strong></p>
    </div>
    
    <div class='content'>
        <p>To,<br/>
        The Hon'ble Judge,<br/>
        [Court Name]</p>
        
        <p>The State of [State] through its Investigating Officer hereby submits this complaint against the accused under Section 302 of the Indian Penal Code, 1860.</p>
        
        <h2>Facts of the Case:</h2>
        [Include detailed facts]
        
        <h2>Charges:</h2>
        [Detail the charges]
        
        <h2>Evidence:</h2>
        [List the evidence]
        
        <h2>Prayer:</h2>
        [State the prayer]
    </div>
    
    <div class='footer' style='margin-top: 40px;'>
        <p>Date: [Current Date]</p>
        <p>Place: [Location]</p>
        <p style='margin-top: 40px;'>Investigating Officer<br/>[Name]<br/>[Designation]</p>
    </div>
</div>

Use the above structure but replace placeholder text with appropriate content based on the case details and IPC sections provided. Ensure proper legal language and formatting throughout."""
        )
        
        # Create chain using the new pipe syntax
        chain = (
            {"case_details": RunnablePassthrough(), "ipc_sections": RunnablePassthrough()} 
            | prompt 
            | model 
            | StrOutputParser()
        )
        
        # Generate document
        result = chain.invoke({
            "case_details": request.case_details,
            "ipc_sections": ", ".join(request.ipc_sections)
        })
        
        return {"document_content": result}
    except Exception as e:
        logger.error(f"Error generating document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")

@router.post("/export_document")
async def export_document(
    request: EditRequest,
    format: str = "pdf"
) -> FileResponse:
    try:
        if not request.document_content:
            raise HTTPException(status_code=400, detail="Missing document content")
        
        if format not in ["pdf", "docx"]:
            raise HTTPException(status_code=400, detail="Unsupported format. Use 'pdf' or 'docx'.")

        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{format}") as tmp_file:
            tmp_path = tmp_file.name
            
            logger.info(f"Exporting document as {format}")
            
            if format == "pdf":
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
                from bs4 import BeautifulSoup
                
                # Create PDF document
                doc = SimpleDocTemplate(
                    tmp_path,
                    pagesize=letter,
                    rightMargin=72,
                    leftMargin=72,
                    topMargin=72,
                    bottomMargin=72
                )
                
                # Create styles
                styles = getSampleStyleSheet()
                styles.add(ParagraphStyle(
                    name='Center',
                    parent=styles['Normal'],
                    alignment=TA_CENTER,
                ))
                
                # Parse HTML content
                soup = BeautifulSoup(request.document_content, 'html.parser')
                story = []
                
                for element in soup.find_all(['h1', 'h2', 'p', 'div']):
                    text = element.get_text().strip()
                    if text:
                        style = 'Center' if element.get('style') and 'text-align: center' in element.get('style') else 'Normal'
                        story.append(Paragraph(text, styles[style]))
                        story.append(Spacer(1, 12))
                
                # Build PDF
                doc.build(story)
                
            else:  # docx
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from bs4 import BeautifulSoup
                
                doc = Document()
                soup = BeautifulSoup(request.document_content, 'html.parser')
                
                for element in soup.find_all(['h1', 'h2', 'p', 'div']):
                    text = element.get_text().strip()
                    if text:
                        p = doc.add_paragraph()
                        p.add_run(text)
                        if element.get('style'):
                            if 'text-align: center' in element.get('style'):
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.save(tmp_path)

            # Return file
            return FileResponse(
                path=tmp_path,
                filename=f"legal_document.{format}",
                media_type=f"application/{'pdf' if format == 'pdf' else 'vnd.openxmlformats-officedocument.wordprocessingml.document'}"
            )

    except Exception as e:
        logger.error(f"Error exporting document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))